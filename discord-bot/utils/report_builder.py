"""Report generation — shared by cogs/report.py and storage/scheduler.py.

build_report_docx: sync, run in executor
generate_ai_summary: async, call before executor
"""

import io
import logging
import os
from collections import Counter, defaultdict
from datetime import datetime, timezone
from typing import List, Optional

logger = logging.getLogger(__name__)


# helpers

def _hex_to_rgb(h: str):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _set_cell_bg(cell, r: int, g: int, b: int):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{r:02X}{g:02X}{b:02X}')
    tcPr.append(shd)


def _style_cell(cell, text: str, bold=False, color=None, size_pt=10, center=False):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    cell.text = ''
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)



async def generate_ai_summary(
    events: List[dict],
    ep_records: List[dict],
    days: int,
    prev_events: Optional[List[dict]] = None,
) -> Optional[str]:
    """call claude for an executive summary, returns None on failure"""
    api_key = os.getenv('ANTHROPIC_API_KEY')
    if not api_key:
        return None
    try:
        import anthropic
        client = anthropic.AsyncAnthropic(api_key=api_key)

        total_ep       = sum(e['ep_awarded'] * len(e['participants']) for e in events)
        unique_members = len({p for e in events for p in e['participants']})
        type_counts    = defaultdict(int)
        for e in events:
            type_counts[e['event_type']] += 1
        unit = os.getenv('REPORT_UNIT_NAME', '327th Star Corps')

        # only send aggregate numbers to Claude — no usernames or member IDs
        top5_ep_values = sorted([r['ep'] for r in ep_records], reverse=True)[:5]
        top5_str = ', '.join(str(ep) for ep in top5_ep_values) if top5_ep_values else 'none'

        prev_context = ""
        if prev_events:
            prev_ep      = sum(e['ep_awarded'] * len(e['participants']) for e in prev_events)
            prev_members = len({p for e in prev_events for p in e['participants']})
            prev_context = (
                f"\nPrevious week comparison:\n"
                f"  Events last week: {len(prev_events)}\n"
                f"  EP last week: {prev_ep}\n"
                f"  Active members last week: {prev_members}\n"
            )

        prompt = (
            f"You are writing the executive summary for a weekly activity report "
            f"of a Roblox military unit called '{unit}'.\n\n"
            f"Report period: {days} days\n"
            f"Total events logged: {len(events)}\n"
            f"Total EP awarded: {total_ep}\n"
            f"Unique active members: {unique_members}\n"
            f"Events by type: {dict(type_counts)}\n"
            f"Top 5 EP totals (highest earners, no names): {top5_str}\n"
            f"{prev_context}\n"
            "Write a professional 2–3 paragraph executive summary. "
            "Be concise and factual. Highlight trends, participation levels, "
            "and overall unit health. Compare to the previous week if data is available. "
            "Paragraph prose only — no bullet points or headers."
        )

        response = await client.messages.create(
            model=os.getenv('CLAUDE_MODEL', 'claude-sonnet-4-6'),
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}],
        )
        return response.content[0].text.strip()
    except Exception as e:
        logger.error(f"AI summary generation failed: {e}")
        return None



def build_report_docx(
    title: str,
    period_str: str,
    events: List[dict],
    ep_records: List[dict],
    ai_summary: Optional[str],
    config: dict,
    prev_events: Optional[List[dict]] = None,
) -> bytes:
    """build the word doc and return raw bytes — must run in executor (sync)"""
    import matplotlib.pyplot as plt
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Inches, Pt, RGBColor

    NAVY   = _hex_to_rgb(config.get('color_primary', '1B2A4A'))
    GOLD   = _hex_to_rgb(config.get('color_accent',  'C9A84C'))
    STRIPE = _hex_to_rgb('EEF2F7')
    WHITE  = (255, 255, 255)
    DARK   = (26, 26, 26)

    unit_name = config.get('unit_name', '327th Star Corps')
    top_n     = int(config.get('top_ep_count', 10))
    days      = int(config.get('days', 7))

    total_events   = len(events)
    total_ep       = sum(e['ep_awarded'] * len(e['participants']) for e in events)
    unique_members = len({p for e in events for p in e['participants']})
    avg_ep_event   = round(total_ep / total_events, 1) if total_events else 0

    doc = Document()
    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    def add_section_heading(text: str):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*NAVY)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), f'{NAVY[0]:02X}{NAVY[1]:02X}{NAVY[2]:02X}')
        pBdr.append(bottom)
        pPr.append(pBdr)
        doc.add_paragraph()

    # cover page
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(unit_name.upper())
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 42)
    run.font.color.rgb = RGBColor(*GOLD)
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(20)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(period_str)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Generated {datetime.now(timezone.utc).strftime('%d %B %Y at %H:%M UTC')}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(130, 130, 130)

    doc.add_paragraph()
    doc.add_paragraph()

    # Metrics table
    if prev_events is not None:
        prev_ep       = sum(e['ep_awarded'] * len(e['participants']) for e in prev_events)
        prev_members  = len({p for e in prev_events for p in e['participants']})
        prev_total    = len(prev_events)
        prev_avg      = round(prev_ep / prev_total, 1) if prev_total else 0

        def _delta(cur, prev):
            d = cur - prev
            return (f'▲ +{d}' if d > 0 else f'▼ {d}' if d < 0 else '— 0') + ' vs last week'

        deltas = [
            _delta(total_events,   prev_total),
            _delta(total_ep,       prev_ep),
            _delta(unique_members, prev_members),
            _delta(avg_ep_event,   prev_avg),
        ]
    else:
        deltas = ['N/A'] * 4

    GREEN = (34, 139, 34)
    RED   = (180, 30, 30)

    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = 'Table Grid'
    labels = ['Total Events', 'EP Awarded', 'Active Members', 'Avg EP / Event']
    values = [str(total_events), str(total_ep), str(unique_members), str(avg_ep_event)]
    for col, (label, value, delta) in enumerate(zip(labels, values, deltas)):
        _set_cell_bg(tbl.rows[0].cells[col], *NAVY)
        _style_cell(tbl.rows[0].cells[col], label,  bold=True, color=WHITE, size_pt=9,  center=True)
        _set_cell_bg(tbl.rows[1].cells[col], *STRIPE)
        _style_cell(tbl.rows[1].cells[col], value,  bold=True, color=DARK,  size_pt=18, center=True)
        _set_cell_bg(tbl.rows[2].cells[col], *WHITE)
        delta_color = GREEN if delta.startswith('▲') else RED if delta.startswith('▼') else (120, 120, 120)
        _style_cell(tbl.rows[2].cells[col], delta, color=delta_color, size_pt=8, center=True)

    doc.add_page_break()

    # section 1: executive summary
    add_section_heading('1.  Executive Summary')
    if ai_summary:
        doc.add_paragraph(ai_summary).style.font.size = Pt(11)
    elif total_events == 0:
        doc.add_paragraph(
            f"No events were logged during this {days}-day reporting period. "
            "Unit activity appears inactive or events have not been recorded."
        )
    else:
        doc.add_paragraph(
            f"During this {days}-day period, the {unit_name} logged {total_events} event(s), "
            f"awarding {total_ep} Engagement Points across {unique_members} unique participant(s). "
            f"On average, {avg_ep_event} EP was awarded per event."
        )
    doc.add_page_break()

    # section 2: event log
    add_section_heading(f'2.  Event Log  ({total_events} events)')
    if not events:
        doc.add_paragraph('No events were logged during this period.')
    else:
        sorted_events = sorted(events, key=lambda e: e['logged_at'])
        tbl2 = doc.add_table(rows=1 + len(sorted_events), cols=5)
        tbl2.style = 'Table Grid'
        for col, h in enumerate(['Date (UTC)', 'Event Type', 'Host', 'Attendees', 'EP Total']):
            _set_cell_bg(tbl2.rows[0].cells[col], *NAVY)
            _style_cell(tbl2.rows[0].cells[col], h, bold=True, color=WHITE, size_pt=9, center=True)
        for i, entry in enumerate(sorted_events, 1):
            dt  = datetime.fromisoformat(entry['logged_at'])
            bg  = STRIPE if i % 2 == 0 else WHITE
            row = tbl2.rows[i]
            data = [
                dt.strftime('%d %b %Y  %H:%M'),
                entry['event_type'],
                entry.get('host_discord_name', '—'),
                str(len(entry['participants'])),
                str(entry['ep_awarded'] * len(entry['participants'])),
            ]
            for col, text in enumerate(data):
                _set_cell_bg(row.cells[col], *bg)
                _style_cell(row.cells[col], text, size_pt=9, center=(col in (3, 4)))
    doc.add_page_break()

    # section 3: charts
    add_section_heading('3.  Activity Overview')
    plt_primary = f'#{NAVY[0]:02X}{NAVY[1]:02X}{NAVY[2]:02X}'
    plt_accent  = f'#{GOLD[0]:02X}{GOLD[1]:02X}{GOLD[2]:02X}'

    if events:
        # bar: events by type
        type_counts: dict = defaultdict(int)
        for e in events:
            type_counts[e['event_type']] += 1

        fig, ax = plt.subplots(figsize=(8, 4))
        types  = list(type_counts.keys())
        counts = [type_counts[t] for t in types]
        bars   = ax.bar(types, counts, color=plt_primary, edgecolor=plt_accent, linewidth=1.5)
        for bar, cnt in zip(bars, counts):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.05,
                    str(cnt), ha='center', va='bottom', fontweight='bold', fontsize=10)
        ax.set_title('Events by Type', fontsize=13, fontweight='bold', pad=12)
        ax.set_xlabel('Event Type', fontsize=10)
        ax.set_ylabel('Count', fontsize=10)
        ax.set_facecolor('#F7F9FC')
        fig.patch.set_facecolor('white')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()
        buf_a = io.BytesIO()
        plt.savefig(buf_a, format='png', dpi=150, bbox_inches='tight')
        buf_a.seek(0)
        plt.close(fig)
        doc.add_picture(buf_a, width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # line: daily events
        daily: Counter = Counter()
        for e in events:
            day = datetime.fromisoformat(e['logged_at']).strftime('%d %b')
            daily[day] += 1

        if len(daily) > 1:
            days_sorted = sorted(daily.keys(), key=lambda d: datetime.strptime(d, '%d %b'))
            counts2 = [daily[d] for d in days_sorted]
            xs = range(len(days_sorted))

            fig2, ax2 = plt.subplots(figsize=(8, 3.5))
            ax2.plot(xs, counts2, marker='o', color=plt_primary, linewidth=2.5,
                     markersize=8, markerfacecolor=plt_accent,
                     markeredgecolor=plt_primary, markeredgewidth=1.5)
            ax2.fill_between(xs, counts2, alpha=0.12, color=plt_primary)
            ax2.set_xticks(list(xs))
            ax2.set_xticklabels(days_sorted, rotation=30, ha='right', fontsize=9)
            ax2.set_title('Daily Event Activity', fontsize=13, fontweight='bold', pad=12)
            ax2.set_ylabel('Events', fontsize=10)
            ax2.set_facecolor('#F7F9FC')
            fig2.patch.set_facecolor('white')
            ax2.spines['top'].set_visible(False)
            ax2.spines['right'].set_visible(False)
            plt.tight_layout()
            buf_b = io.BytesIO()
            plt.savefig(buf_b, format='png', dpi=150, bbox_inches='tight')
            buf_b.seek(0)
            plt.close(fig2)
            doc.add_picture(buf_b, width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph('No event data available for charting.')
    doc.add_page_break()

    # section 4: leaderboard
    add_section_heading(f'4.  EP Leaderboard  (Top {top_n})')
    top_members = sorted(ep_records, key=lambda r: r['ep'], reverse=True)[:top_n]
    if top_members:
        tbl3 = doc.add_table(rows=1 + len(top_members), cols=3)
        tbl3.style = 'Table Grid'
        for col, h in enumerate(['Rank', 'Roblox Username', 'EP']):
            _set_cell_bg(tbl3.rows[0].cells[col], *NAVY)
            _style_cell(tbl3.rows[0].cells[col], h, bold=True, color=WHITE, size_pt=10, center=True)
        for rank, record in enumerate(top_members, 1):
            bg    = STRIPE if rank % 2 == 0 else WHITE
            row   = tbl3.rows[rank]
            medal = {1: '🥇', 2: '🥈', 3: '🥉'}.get(rank, f'#{rank}')
            for col, text in enumerate([medal, record['roblox_username'], str(record['ep'])]):
                _set_cell_bg(row.cells[col], *bg)
                _style_cell(row.cells[col], text, size_pt=10, center=(col != 1))
    else:
        doc.add_paragraph('No EP records found.')

    # closing note
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"{unit_name}  •  Auto-generated report  •  "
        f"{datetime.now(timezone.utc).strftime('%d %B %Y')}"
    )
    run.font.size   = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(160, 160, 160)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
